"""
20_generate_report_utt.py — Rapport 10 pages, thème UTT professionnel.
Page de garde (bandeau + logo), sommaire automatique (TOC), pagination,
en-tête avec logo, police Segoe UI, couleurs UTT. Contenu = données réelles
ALT1/ALT2 + régression. Produit reports/Rapport_IF23_ALT3_10pages.docx.
"""
from __future__ import annotations

import sys
from pathlib import Path

try:
    sys.stdout.reconfigure(encoding="utf-8")
except Exception:
    pass

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm, Inches

ROOT = Path(__file__).resolve().parents[1]
FIG = ROOT / "reports" / "figures"
ASSETS = ROOT / "reports" / "assets"
OUT = ROOT / "reports" / "Rapport_IF23_ALT3_10pages.docx"

NAVY = RGBColor(0x00, 0x1E, 0x62); BLUE = RGBColor(0x00, 0x72, 0xCE)
CYAN = RGBColor(0x00, 0x9C, 0xDE); GREY = RGBColor(0x55, 0x55, 0x55); RED = RGBColor(0xC0, 0x10, 0x2E)
FONT = "Segoe UI"


def csv(n): return pd.read_csv(FIG / n)


# ----------------------------------------------------------------- low-level
def _field(run, instr, placeholder=""):
    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
    i = OxmlElement("w:instrText"); i.set(qn("xml:space"), "preserve"); i.text = instr
    s = OxmlElement("w:fldChar"); s.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t"); t.text = placeholder
    e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end")
    for el in (b, i, s, t, e):
        run._r.append(el)


def setfont(doc):
    st = doc.styles["Normal"]; st.font.name = FONT; st.font.size = Pt(10.5)
    st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    for lvl, sz in [("Heading 1", 14), ("Heading 2", 11.5), ("Heading 3", 10.5)]:
        s = doc.styles[lvl]; s.font.name = FONT; s.font.size = Pt(sz)
        s.font.color.rgb = NAVY; s.font.bold = True


def update_fields_on_open(doc):
    el = OxmlElement("w:updateFields"); el.set(qn("w:val"), "true")
    doc.settings.element.append(el)


def _flow(pf, together=True, with_next=False):
    """Évite les paragraphes coupés en deux et les lignes orphelines."""
    pf.widow_control = True
    pf.keep_together = together      # le paragraphe ne se coupe pas entre deux pages
    if with_next:
        pf.keep_with_next = True     # reste collé à l'élément suivant (titre, légende)


def h(doc, t, lvl=1):
    p = doc.add_heading(t, level=lvl)
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(3)
    _flow(p.paragraph_format, together=True, with_next=True)  # titre jamais seul en bas de page
    return p


def para(doc, t, size=10.5, italic=False, bold=False, color=None, align=None, after=4):
    p = doc.add_paragraph(); r = p.add_run(t)
    r.italic = italic; r.bold = bold; r.font.size = Pt(size); r.font.name = FONT
    if color: r.font.color.rgb = color
    if align: p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    _flow(p.paragraph_format)
    return p


def bullet(doc, t, pre=None):
    p = doc.add_paragraph(style="List Bullet")
    if pre:
        r = p.add_run(pre); r.bold = True; r.font.name = FONT; r.font.size = Pt(10)
    r2 = p.add_run(t); r2.font.name = FONT; r2.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(1)
    _flow(p.paragraph_format)
    return p


def fig(doc, name, cap, w=5.4):
    if not (FIG / name).exists():
        para(doc, f"[figure manquante: {name}]", italic=True, color=RED); return
    doc.add_picture(str(FIG / name), width=Inches(w))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.paragraphs[-1].paragraph_format.space_after = Pt(1)
    _flow(doc.paragraphs[-1].paragraph_format, together=True, with_next=True)  # image collée à sa légende
    c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = c.add_run(cap); r.italic = True; r.font.size = Pt(8.5); r.font.name = FONT; r.font.color.rgb = GREY
    c.paragraph_format.space_after = Pt(8); _flow(c.paragraph_format)


def shade(cell, hx):
    tcPr = cell._tc.get_or_add_tcPr(); s = OxmlElement("w:shd")
    s.set(qn("w:val"), "clear"); s.set(qn("w:fill"), hx); tcPr.append(s)


def _no_borders(t):
    tblPr = t._tbl.tblPr; b = OxmlElement("w:tblBorders")
    for e in ("top", "left", "bottom", "right", "insideH", "insideV"):
        x = OxmlElement(f"w:{e}"); x.set(qn("w:val"), "none"); b.append(x)
    tblPr.append(b)


def _cant_split(t):
    for row in t.rows:
        trPr = row._tr.get_or_add_trPr()
        cs = OxmlElement("w:cantSplit"); cs.set(qn("w:val"), "true"); trPr.append(cs)


def two_figs(doc, f1, cap1, f2, cap2, w1=3.0, w2=3.0):
    """Deux figures côte à côte dans une table sans bordures."""
    t = doc.add_table(rows=1, cols=2); _no_borders(t); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    _cant_split(t)
    for cell, fn, cap, w in [(t.rows[0].cells[0], f1, cap1, w1), (t.rows[0].cells[1], f2, cap2, w2)]:
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if (FIG / fn).exists():
            p.add_run().add_picture(str(FIG / fn), width=Inches(w))
        cp = cell.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cp.add_run(cap); r.italic = True; r.font.size = Pt(8); r.font.name = FONT; r.font.color.rgb = GREY


def table(doc, df, headers, pct=None, fs=8.5):
    pct = pct or []
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, ht in enumerate(headers):
        c = t.rows[0].cells[j]; c.text = ""
        r = c.paragraphs[0].add_run(str(ht)); r.bold = True; r.font.size = Pt(fs); r.font.name = FONT
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); shade(c, "001E62")
    for k, (_, row) in enumerate(df.iterrows()):
        cs = t.add_row().cells
        for j, col in enumerate(df.columns):
            v = row[col]
            txt = f"{v*100:.1f} %" if (col in pct and isinstance(v, (int, float))) else \
                  (f"{v:.3f}" if isinstance(v, float) else str(v))
            cs[j].text = ""; r = cs[j].paragraphs[0].add_run(txt); r.font.size = Pt(fs); r.font.name = FONT
            if k % 2 == 1:
                shade(cs[j], "EAF2FB")
    _cant_split(t)  # une ligne de tableau ne se coupe pas entre deux pages
    # répéter la ligne d'en-tête si le tableau passe une page
    trPr = t.rows[0]._tr.get_or_add_trPr(); th = OxmlElement("w:tblHeader"); th.set(qn("w:val"), "true"); trPr.append(th)


def note(doc, t):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.3)
    r = p.add_run("➤  " + t); r.italic = True; r.font.size = Pt(9); r.font.name = FONT; r.font.color.rgb = BLUE
    _flow(p.paragraph_format)


def rule(doc):
    p = doc.add_paragraph(); pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr"); bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "12"); bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "009CDE"); pbdr.append(bottom); pPr.append(pbdr)
    p.paragraph_format.space_after = Pt(2)


# ----------------------------------------------------------------- header/footer
def setup_header_footer(section):
    section.different_first_page_header_footer = True
    # En-tête (pages courantes) : logo + titre
    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if (ASSETS / "utt_logo.png").exists():
        hp.add_run().add_picture(str(ASSETS / "utt_logo.png"), width=Inches(0.85))
    run = hp.add_run("      IF23 P — Localisation indoor WiFi")
    run.font.name = FONT; run.font.size = Pt(8.5); run.font.color.rgb = GREY
    # Pied de page : Page X / Y
    fp = section.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = fp.add_run("Page "); r1.font.size = Pt(8.5); r1.font.name = FONT; r1.font.color.rgb = GREY
    rp = fp.add_run(); rp.font.size = Pt(8.5); rp.font.name = FONT; rp.font.color.rgb = GREY
    _field(rp, "PAGE", "1")
    r2 = fp.add_run(" / "); r2.font.size = Pt(8.5); r2.font.name = FONT; r2.font.color.rgb = GREY
    rn = fp.add_run(); rn.font.size = Pt(8.5); rn.font.name = FONT; rn.font.color.rgb = GREY
    _field(rn, "NUMPAGES", "1")


# ----------------------------------------------------------------- pages
def title_page(doc):
    if (ASSETS / "utt_banner.png").exists():
        doc.add_picture(str(ASSETS / "utt_banner.png"), width=Inches(6.7))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.paragraphs[-1].paragraph_format.space_after = Pt(6)
    if (ASSETS / "utt_logo.png").exists():
        doc.add_picture(str(ASSETS / "utt_logo.png"), width=Inches(2.2))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(doc, "UV IF23 P — Géolocalisation indoor et outdoor", size=11, color=GREY,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=18)
    para(doc, "Localisation indoor par fingerprinting WiFi", size=24, bold=True, color=NAVY,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    para(doc, "Détection de salle et estimation de position dans le bâtiment de l'UTT",
         size=13, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    rule(doc)
    para(doc, "Binôme 6 — Molka GHADDAB · Maxime LE GAL · Virgile TUPET",
         size=12, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, "Encadrante : Farah Chehade", size=11, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, "Jalons ALT1 · ALT2 · ALT3 — Semestre P26 (juin 2026)", size=11, color=GREY,
         align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()


def sommaire(doc):
    para(doc, "Sommaire", size=16, bold=True, color=NAVY, after=6)
    p = doc.add_paragraph(); r = p.add_run()
    _field(r, 'TOC \\o "1-2" \\h \\z \\u',
           "Sommaire : clic droit puis « Mettre à jour les champs » si non affiché.")
    doc.add_page_break()


# ----------------------------------------------------------------- content
def content(doc):
    head = csv("honest_headline_bssid.csv").iloc[0]

    # ---- 1
    h(doc, "1. Introduction et objectifs")
    h(doc, "1.1 Contexte", 2)
    para(doc,
        "À l'intérieur d'un bâtiment, le GPS est inutilisable : les murs atténuent fortement le "
        "signal et la notion d'étage disparaît. Les réseaux WiFi, eux, sont présents partout et "
        "émettent en permanence. On peut donc s'en servir pour se localiser, sans installer de "
        "matériel supplémentaire. L'idée est le fingerprinting : à chaque endroit, un appareil "
        "capte un ensemble de bornes WiFi avec une puissance (RSSI, en dBm) qui lui est propre. "
        "Cet ensemble de puissances forme une « empreinte » radio caractéristique du lieu. On "
        "constitue d'abord une base associant des empreintes à des positions (phase hors ligne), "
        "puis on compare une nouvelle mesure à cette base pour estimer la position (phase en ligne).")
    h(doc, "1.2 Objectifs et cahier des charges", 2)
    para(doc,
        "Conformément à la consigne, le projet traite deux problèmes complémentaires : prédire la "
        "zone (salle) où se trouve l'appareil — un problème de classification — et estimer ses "
        "coordonnées exactes (x, y) à l'intérieur d'une zone — un problème de régression. À cela "
        "s'ajoute une exigence centrale : mesurer comment les performances évoluent dans le temps, "
        "à travers trois campagnes de mesures successives (ALT1, ALT2, ALT3).")
    spec = pd.DataFrame({
        "Exigence": ["Classification", "Régression", "Vectorisation", "Modèles",
                     "Évolution temporelle", "Outils"],
        "Description": ["prédire la zone parmi N (salle / couloir)",
                        "estimer (x, y) dans une zone donnée",
                        "vecteur (RSSI BSSID₁ … BSSIDₙ), -100 si non capté",
                        "plusieurs algorithmes, avec tous les réseaux puis une sélection",
                        "stratégies ALT1 / ALT2 / ALT3 (obligatoire)",
                        "Python, pywifi, pandas / scikit-learn"]})
    table(doc, spec, list(spec.columns))
    para(doc, "Tableau 1 — Synthèse du cahier des charges.", italic=True, size=8.5, color=GREY)
    para(doc,
        "Le périmètre couvre les bâtiments P et S de l'UTT sur plusieurs étages. Selon le découpage "
        "retenu, le système traite de 10 à 24 zones (salles, couloirs, paliers).")

    # ---- 2
    h(doc, "2. Données et collecte")
    h(doc, "2.1 Nature des mesures", 2)
    para(doc,
        "Chaque relevé associe le nom du réseau (SSID), son adresse matérielle (BSSID), sa "
        "puissance (RSSI) et l'instant de mesure. À chaque point, on enchaîne plusieurs scans que "
        "l'on moyenne par réseau, ce qui lisse le bruit radio. On obtient un vecteur de puissances, "
        "une valeur par réseau connu ; les réseaux non détectés à ce point reçoivent -100 dBm "
        "(signal absent). La base de données est organisée par bâtiment, étage, zone, coordonnées "
        "(x, y), réseau et instant de mesure.")
    h(doc, "2.2 Les trois campagnes ALT1, ALT2, ALT3", 2)
    para(doc,
        "La consigne définit trois campagnes successives, de densité décroissante, dont l'objet est "
        "de mesurer la tenue du modèle dans le temps. ALT1 (février) est la campagne de référence, "
        "la plus dense ; ALT2 (avril) re-mesure une partie des zones trois mois plus tard ; ALT3 "
        "(juin) servira de test final. Le protocole a évolué entre ALT1 (sous-découpage de chaque "
        "salle en deux sous-zones A/B) et ALT2 (capture continue de la salle entière, environ trois "
        "minutes par salle). Pour comparer les deux campagnes, les sous-zones A/B d'ALT1 sont "
        "fusionnées au niveau de la salle.")
    diff = pd.DataFrame({
        "Jeu": ["ALT1", "ALT2", "ALT3"],
        "Période": ["février 2026", "28 avril 2026", "juin 2026 (à venir)"],
        "Zones": ["24 (sous-zones A/B)", "10 (salles entières)", "—"],
        "Snapshots": ["1 498", "452", "—"],
        "Rôle": ["apprentissage", "test temporel + ré-app.", "test final (stratégie 3)"]})
    table(doc, diff, list(diff.columns))
    para(doc, "Tableau 2 — Les trois campagnes. ALT1 et ALT2 sont intégrées ; ALT3 reste à "
              "collecter.", italic=True, size=8.5, color=GREY)
    para(doc,
        "Sept zones sont communes à ALT1 et ALT2 (elles permettent la comparaison temporelle) et "
        "trois sont nouvelles en ALT2 (P104, P204, S201). À noter, un point de rigueur sur les "
        "données : un fichier ALT2 contenait en réalité deux salles capturées à la suite ; un écart "
        "temporel d'environ 80 s a permis de les séparer automatiquement en deux fichiers distincts.")
    h(doc, "2.3 Jeu de données pour la régression", 2)
    para(doc,
        "Pour la position exacte, une salle (environ 5 × 7 m) a été cartographiée finement : 46 "
        "points de coordonnées (x, y) connues, chacun décrit par le RSSI de 52 bornes. Ce jeu, "
        "conforme au format attendu (X, Y puis une colonne par BSSID), sert à l'apprentissage de la "
        "régression intra-zone (section 6).")

    # ---- 3
    h(doc, "3. Ingénierie des caractéristiques")
    h(doc, "3.1 Vectorisation des scans", 2)
    para(doc,
        "Le défi central est de transformer un scan WiFi — une liste de longueur variable de "
        "triplets (SSID, BSSID, RSSI) — en un vecteur de taille fixe, identique pour tous les "
        "relevés. Pour chaque réseau retenu (présent dans au moins quelques snapshots), on calcule "
        "des statistiques de puissance : moyenne, maximum, écart-type, et présence. On y ajoute des "
        "indicateurs globaux du scan (nombre de réseaux visibles, RSSI moyen, médian, etc.). Cette "
        "transformation est encapsulée dans un module réutilisable, garantissant la même opération "
        "à l'apprentissage et à la prédiction.")
    h(doc, "3.2 Choix de l'identifiant réseau", 2)
    para(doc,
        "Le choix de l'identifiant a un effet direct sur les performances. Trois options ont été "
        "comparées : le SSID seul (nom de réseau), le BSSID seul (borne physique), et un identifiant "
        "composé SSID|BSSID qui combine les deux. Le SSID est trop peu discriminant (tous les points "
        "d'accès UTT partagent les mêmes noms) ; le BSSID distingue chaque borne ; l'identifiant "
        "composé apporte le maximum d'information.")
    enc = pd.DataFrame({
        "Identifiant": ["SSID seul", "BSSID seul", "SSID|BSSID (composé)"],
        "Nb identifiants": ["13", "104", "139"],
        "Accuracy hold-out": ["98,3 %", "99,0 %", "100,0 %"],
        "Remarque": ["peu discriminant", "discriminant", "le plus riche"]})
    table(doc, enc, list(enc.columns))
    para(doc, "Tableau 3 — Effet de l'identifiant réseau (accuracy sur hold-out aléatoire).",
         italic=True, size=8.5, color=GREY)
    h(doc, "3.3 Couverture et corrélation des réseaux", 2)
    para(doc,
        "Tous les réseaux ne se valent pas. La couverture est très inégale (figure 1) : quelques "
        "réseaux institutionnels sont présents presque partout, tandis qu'une longue traîne de "
        "réseaux n'apparaît que ponctuellement. De plus, beaucoup sont redondants : on relève 285 "
        "paires de réseaux très corrélées (|r| > 0,9), principalement les réseaux UTT et eduroam "
        "portés par les mêmes bornes physiques. Cette redondance justifie une sélection de réseaux "
        "plutôt que d'utiliser l'ensemble brut.")
    fig(doc, "fig17_couverture.png",
        "Figure 1 — Couverture des réseaux : quelques stables, beaucoup d'éphémères.", w=4.0)

    # ---- 4
    h(doc, "4. Classification en zones")
    h(doc, "4.1 Modèles comparés", 2)
    para(doc,
        "Plusieurs familles de modèles vues en cours ont été évaluées, d'abord avec tous les réseaux "
        "puis avec une sélection : k plus proches voisins, classifieur bayésien, SVM à noyau RBF, "
        "forêts aléatoires et ExtraTrees, et réseaux de neurones. Les méthodes à base d'arbres "
        "arrivent en tête ; un vote combinant forêt aléatoire, ExtraTrees et régression logistique "
        "atteint les meilleurs scores sur le hold-out.")
    h(doc, "4.2 Méthodologie d'évaluation", 2)
    para(doc,
        "La manière d'évaluer change radicalement les conclusions. Les snapshots d'une même zone "
        "proviennent d'une seule session de mesure et se ressemblent beaucoup d'un instant au "
        "suivant. Un découpage aléatoire entre apprentissage et test place donc des mesures presque "
        "identiques des deux côtés : le score obtenu est optimiste. Pour s'en approcher des "
        "conditions réelles, nous durcissons progressivement le protocole — séparation des données "
        "dans le temps, puis test sur une campagne entièrement nouvelle (ALT2). La figure 2 montre "
        "cette descente, du test de base jusqu'au test réel sur ALT2.")
    fig(doc, "fig13_escalier.png",
        "Figure 2 — Du test aléatoire (99 %) au test réel sur ALT2 (47 %), puis récupération.", w=5.0)
    sb = csv("eval_summary_bssid.csv")
    table(doc, sb[["model", "A_aleatoire", "C2_groupkfold"]],
          ["Modèle", "Test aléatoire", "Évaluation réaliste"], pct=["A_aleatoire", "C2_groupkfold"])
    para(doc, "Tableau 4 — Accuracy par modèle : test aléatoire vs évaluation réaliste (séparation "
              "temporelle, BSSID).", italic=True, size=8.5, color=GREY)
    h(doc, "4.3 Analyse des résultats", 2)
    para(doc,
        f"Avec un protocole plus exigeant, l'accuracy se situe autour de 80 % en validation "
        f"temporelle, ce qui reste cohérent avec la littérature pour de la classification par pièce. "
        f"Par ailleurs, même dans ce cadre, la bonne zone figure dans les trois premières "
        f"propositions dans {head['top3']*100:.1f} % des cas : le système est donc très fiable pour "
        "proposer une courte liste de zones candidates. L'analyse de l'importance des réseaux "
        "(figure 3) apporte un éclairage utile : les deux réseaux les plus discriminants sont des "
        "téléphones personnels présents lors de la collecte. Ils aident sur le moment, mais ne "
        "seront pas là d'une campagne à l'autre, ce qui annonce la dégradation observée en "
        "section 5.")
    fig(doc, "fig4_importance_reseaux.png",
        "Figure 3 — Réseaux les plus discriminants. Les deux premiers sont des téléphones personnels.", w=4.6)

    # ---- 5
    h(doc, "5. Évolution temporelle (ALT1 → ALT2 → ALT3)")
    ev = csv("evolution_real.csv")
    s1 = float(ev.loc[ev["strategie"].str.startswith("S1"), "accuracy"].iloc[0])
    s2 = float(ev.loc[ev["strategie"].str.startswith("S2"), "accuracy"].iloc[0])
    para(doc,
        "C'est le résultat central du projet. On vérifie si un modèle appris en février reconnaît "
        "encore des mesures prises en avril (ALT2), trois mois plus tard. Trois stratégies sont "
        "définies : la stratégie 1 entraîne sur ALT1 et teste sur ALT2 (mesure du drift pur) ; la "
        "stratégie 2 ajoute la moitié d'ALT2 à l'apprentissage et teste sur le reste (effet d'un "
        "ré-entraînement partiel) ; la stratégie 3 entraînera sur ALT1 + ALT2 et testera sur ALT3.")
    para(doc,
        f"En stratégie 1, l'accuracy descend à {s1*100:.0f} % : le modèle de février ne reconnaît "
        "plus la moitié des mesures d'avril, et plusieurs zones tombent à zéro (figure 4), leur "
        "environnement radio ayant changé (bornes déplacées ou remplacées, réseaux invités "
        "différents). En réintégrant seulement la moitié d'ALT2 à l'apprentissage (stratégie 2), on "
        f"remonte à {s2*100:.0f} %. Le contraste entre ces deux chiffres quantifie le « drift » du "
        "WiFi : l'empreinte radio d'un lieu n'est pas figée, et un système de fingerprinting doit "
        "être ré-entraîné régulièrement pour rester fiable — un enseignement directement exploitable "
        "pour un déploiement réel.")
    fig(doc, "fig18_evolution.png",
        "Figure 4 — Évolution ALT1 → ALT2 : accuracy par zone (gauche), drift puis récupération (droite).", w=5.5)
    note(doc, "Stratégie 3 (test sur ALT3) : en attente de la collecte ALT3 (juin). Le code est prêt "
              "et le tableau se complétera automatiquement.")

    # ---- 6
    h(doc, "6. Régression intra-zone (position exacte)")
    rr = csv("regression_final.csv").sort_values("MAE_m"); best = rr.iloc[0]
    para(doc,
        "Pour estimer la position (x, y) à l'intérieur d'une zone, on apprend une régression à "
        "partir des 46 points cartographiés. La qualité se mesure par l'erreur euclidienne (moyenne, "
        "médiane, 90ᵉ percentile) et par la courbe cumulative des erreurs (CDF). Nous avons constaté "
        "que les deux axes n'ont pas la même difficulté : l'axe Y est bien prédit par un modèle "
        "d'arbres (ExtraTrees), tandis que l'axe X, plus délicat, est mieux capté par un processus "
        "gaussien. Le meilleur résultat vient donc d'un modèle hybride combinant les deux, un par "
        "axe.")
    table(doc, rr.rename(columns={"modele": "Modèle", "MAE_m": "MAE (m)", "med_m": "Médiane (m)",
          "p90_m": "P90 (m)", "R2_x": "R² (X)", "R2_y": "R² (Y)"}),
          ["Modèle", "MAE (m)", "Médiane (m)", "P90 (m)", "R² (X)", "R² (Y)"])
    para(doc, "Tableau 5 — Erreur de position par modèle (validation Leave-One-Out).",
         italic=True, size=8.5, color=GREY)
    para(doc,
        f"Le modèle hybride atteint une erreur médiane de {best['med_m']:.2f} m et une moyenne de "
        f"{best['MAE_m']:.2f} m (90 % des points sous {best['p90_m']:.1f} m), soit une position estimée à "
        "environ un mètre près dans la salle. Les figures 5 et 6 montrent la CDF des erreurs et la "
        "comparaison entre positions vraies et prédites.")
    two_figs(doc, "fig19_regression_cdf.png", "Figure 5 — CDF de l'erreur de position.",
             "fig20_regression_scatter.png", "Figure 6 — Positions vraies vs prédites.",
             w1=3.1, w2=2.7)
    note(doc, "La précision est aujourd'hui limitée surtout par le nombre de points (46) et la "
              "taille de la zone : le levier principal pour gagner en précision est une grille de "
              "mesure plus dense, plus que le choix du modèle.")

    # ---- 7
    h(doc, "7. Évolutions avancées")
    by = csv("bayesian_bssid.csv"); bc = by[by["modele"].str.contains("cours")].iloc[0]
    tu = csv("hyperparam_tuning_bssid.csv"); etg = tu[tu["modele"] == "ExtraTrees"].iloc[0]
    hh = csv("hierarchical_bssid.csv")
    gf = float(hh.loc[hh["approche"].str.contains("étage"), "accuracy"].iloc[0])
    para(doc, "Plusieurs pistes ont été explorées au-delà du modèle de base :")
    bullet(doc, f"programmé d'après le cours (loi gaussienne du RSSI par réseau, présence modélisée, "
                f"décision au maximum a posteriori). Il atteint {bc['groupkfold_acc']*100:.0f} % en "
                "évaluation réaliste, comparable aux forêts.", "Modèle bayésien — ")
    bullet(doc, f"recherche sur grille avec validation croisée temporelle ; les gains sont modestes "
                f"(ExtraTrees +{etg['gain']*100:.0f} point). Le protocole d'évaluation et le choix "
                "des réseaux pèsent davantage que le réglage fin.", "Optimisation des hyperparamètres — ")
    bullet(doc, f"prédire d'abord l'étage (fiable à {gf*100:.0f} %) puis la salle ; la grande "
                "fiabilité de l'étage est une base solide pour la cartographie.", "Approche hiérarchique — ")
    bullet(doc, "estimateur log-distance (exposant de propagation n ≈ 2,8) reliant puissance et "
                "distance à une borne, utile au diagnostic (bonus hors consigne).", "Estimation de distance — ")
    bullet(doc, "limiter les zones possibles d'un instant à l'autre via un graphe de transition "
                "entre zones et une vitesse maximale. Piste prévue.", "Exploitation de la mobilité — ")
    fig(doc, "fig12_hierarchique.png",
        "Figure 7 — Approche hiérarchique : la détection d'étage (94 %) servira à la cartographie.", w=4.6)

    # ---- 8 Cartographie
    h(doc, "8. Cartographie")
    para(doc,
        "La localisation est restituée sur deux niveaux : le plan du campus UTT situe le bâtiment "
        "(section S ou P), puis un croquis d'étage situe la salle et la position à l'intérieur. "
        "L'agencement des étages suit la réalité du bâtiment : par section et par étage, deux salles "
        "côte à côte d'un côté du couloir, deux autres en face, et un palier qui relie les niveaux.")
    h(doc, "8.1 Cartographie en puissance", 2)
    para(doc,
        "À partir des points de la zone cartographiée, on interpole la puissance reçue de chaque "
        "borne sur toute la surface (figure 8). Ces cartes de chaleur montrent la structure spatiale "
        "du signal : elle est nette selon Y mais faible selon X, ce qui explique pourquoi l'axe X "
        "est plus difficile à prédire en régression (section 6).")
    fig(doc, "fig21_heatmaps.png",
        "Figure 8 — Cartographie en puissance : RSSI interpolé de quatre bornes sur la zone "
        "(points blancs = mesures).", w=4.6)
    h(doc, "8.2 Localisation de bout en bout", 2)
    para(doc,
        "La chaîne complète est illustrée sur un scan réel de la campagne ALT2 (figure 9) : le "
        "système identifie l'étage et la section, prédit la salle avec son niveau de confiance et le "
        "top-3 des zones probables, puis place la position estimée (x, y) à l'intérieur de la salle "
        "sur le croquis.")
    fig(doc, "fig24_demo_localisation.png",
        "Figure 9 — Localisation de bout en bout sur un scan réel : zone prédite sur le croquis "
        "d'étage + top-3 des zones.", w=6.0)
    h(doc, "8.3 Aide à la mobilité (désactivable)", 2)
    ma = csv("map_assist.csv")
    off = float(ma.iloc[0]["accuracy_trajectoire"]); on = float(ma.iloc[1]["accuracy_trajectoire"])
    para(doc,
        "Lorsqu'on suit une trajectoire, on ne peut pas passer instantanément à une zone éloignée. "
        "On peut donc contraindre les prédictions successives par un graphe de transition entre "
        "zones voisines (lissage de la trajectoire). Sur des trajectoires simulées respectant "
        f"l'adjacence, cette aide fait passer l'accuracy de {off*100:.0f} % à {on*100:.0f} % "
        f"(+{(on-off)*100:.0f} points, figure 10). Conformément au principe de prudence retenu, "
        "cette aide reste désactivée par défaut et n'est activée que si elle améliore réellement la "
        "prédiction.")
    fig(doc, "fig23_map_assist.png",
        "Figure 10 — Aide à la mobilité : impact mesuré (désactivable, désactivée par défaut).", w=4.3)

    # ---- 9
    h(doc, "9. Application temps réel")
    para(doc,
        "Le pipeline a été intégré dans des interfaces graphiques permettant de tester le système en "
        "direct sur un ordinateur, à l'aide de pywifi :")
    bullet(doc, "scanne le WiFi en continu, affiche la salle prédite, sa confiance et le top-5 des "
                "zones probables ; l'utilisateur peut indiquer la salle réelle pour suivre une "
                "accuracy en direct, et enregistrer un scan au format d'entraînement.",
                "Classification de salle — ")
    bullet(doc, "affiche la position (x, y) estimée sur un plan de la zone, avec un indicateur du "
                "nombre de bornes connues effectivement captées.", "Position intra-zone — ")
    bullet(doc, "estime la distance à une borne à partir du modèle log-distance (bonus).",
                "Distance à une borne — ")
    para(doc,
        "Ces interfaces rendent le système concret et serviront de base à la cartographie (afficher "
        "la zone puis la position précise sur un plan du bâtiment).")

    # ---- 10
    h(doc, "10. Difficultés rencontrées")
    bullet(doc, "les modèles récents attendent l'identifiant composé SSID|BSSID, alors que le "
                "scanner live produisait au départ des SSID seuls : tout le vecteur arrivait vide et "
                "la prédiction devenait aléatoire. La correction détecte automatiquement le format "
                "attendu par le modèle chargé.", "Cohérence d'encodage entre apprentissage et live — ")
    bullet(doc, "lors d'un test en direct, aucun des BSSID captés ne figurait dans les bases "
                "d'apprentissage : l'infrastructure WiFi avait été en partie renouvelée (mêmes "
                "fabricants, adresses différentes). C'est une illustration concrète du drift et une "
                "limite intrinsèque du fingerprinting — on ne reconnaît que les empreintes déjà "
                "vues.", "Renouvellement de l'infrastructure — ")
    bullet(doc, "une zone à très peu de mesures a dû être écartée de la validation croisée stratifiée ; "
                "les conflits de version de scikit-learn entre artefacts ont été gérés sans impact "
                "fonctionnel.", "Petites classes et versions — ")

    # ---- 11
    h(doc, "11. Conclusion et perspectives")
    para(doc,
        "Le projet traite les deux tâches demandées : la classification en zones et la régression de "
        "position (x, y), avec une erreur médiane d'environ un mètre. Sur le hold-out, la "
        "classification atteint des scores très élevés, mais l'enseignement principal vient de "
        "l'évolution temporelle : sur une nouvelle campagne (ALT2, trois mois après), l'accuracy "
        "chute à environ 47 % et ne se rétablit qu'en réintégrant des mesures récentes. Le "
        "fingerprinting WiFi fonctionne, mais reste sensible au temps : il faut le ré-entraîner "
        "régulièrement, car l'environnement radio change.")
    para(doc,
        "Les perspectives sont claires : collecter ALT3 (juin) pour compléter la stratégie 3 (le "
        "code est prêt), densifier la régression et l'étendre à d'autres zones, et intégrer la "
        "cartographie (déjà en place : carte de puissance, croquis d'étage, aide à la mobilité) "
        "dans l'interface temps réel. L'apport "
        "méthodologique du projet est d'avoir distingué un score de hold-out optimiste d'une "
        "performance mesurée en conditions réalistes, ce qui donne une image fidèle de ce que le "
        "système sait réellement faire.")


def build():
    doc = Document(); setfont(doc)
    sec = doc.sections[0]
    sec.top_margin = Cm(1.7); sec.bottom_margin = Cm(1.5); sec.left_margin = Cm(1.9); sec.right_margin = Cm(1.9)
    sec.header_distance = Cm(0.8); sec.footer_distance = Cm(0.8)
    setup_header_footer(sec)
    title_page(doc)
    sommaire(doc)
    content(doc)
    update_fields_on_open(doc)
    OUT.parent.mkdir(parents=True, exist_ok=True); doc.save(OUT)
    print("Rapport UTT généré :", OUT)


if __name__ == "__main__":
    build()
