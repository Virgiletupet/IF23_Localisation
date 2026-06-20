"""
15_generate_report_10p.py — Version condensée (~10 pages) recadrée sur la consigne.
ALT1/ALT2/ALT3 = jeux de données successifs. Couvre classification ET régression,
l'évolution temporelle (3 stratégies) et les évolutions avancées.
Produit reports/Rapport_IF23_ALT3_10pages.docx.
"""
from __future__ import annotations

import sys
from pathlib import Path

try:
    sys.stdout.reconfigure(encoding="utf-8")
except Exception:
    pass

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm, Inches

ROOT = Path(__file__).resolve().parents[1]
FIG = ROOT / "reports" / "figures"
OUT = ROOT / "reports" / "Rapport_IF23_ALT3_10pages.docx"
NAVY = RGBColor(0x0B, 0x3D, 0x6B); RED = RGBColor(0xC0, 0x39, 0x2B); GREY = RGBColor(0x55, 0x55, 0x55)


def csv(n): return pd.read_csv(FIG / n)


def setfont(doc):
    st = doc.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(10.5)
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")


def h(doc, t, lvl=1):
    p = doc.add_heading(t, level=lvl)
    for r in p.runs:
        r.font.color.rgb = NAVY if lvl <= 2 else GREY
        r.font.size = Pt(13 if lvl == 1 else 11)
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(2)
    return p


def para(doc, t, size=10.5, italic=False, bold=False, color=None, align=None):
    p = doc.add_paragraph(); r = p.add_run(t)
    r.italic = italic; r.bold = bold; r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    if align: p.alignment = align
    p.paragraph_format.space_after = Pt(3)
    return p


def bullet(doc, t, pre=None):
    p = doc.add_paragraph(style="List Bullet")
    if pre:
        r = p.add_run(pre); r.bold = True
    p.add_run(t)
    for r in p.runs: r.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(1)
    return p


def fig(doc, name, cap, w=5.3):
    if not (FIG / name).exists():
        para(doc, f"[figure manquante: {name}]", italic=True, color=RED); return
    doc.add_picture(str(FIG / name), width=Inches(w))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.paragraphs[-1].paragraph_format.space_after = Pt(1)
    c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = c.add_run(cap); r.italic = True; r.font.size = Pt(8.5); r.font.color.rgb = GREY
    c.paragraph_format.space_after = Pt(4)


def shade(cell, hx):
    tcPr = cell._tc.get_or_add_tcPr(); sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), hx); tcPr.append(sh)


def table(doc, df, headers, pct=None, fs=8.5):
    pct = pct or []
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
    for j, ht in enumerate(headers):
        c = t.rows[0].cells[j]; c.text = ""
        rn = c.paragraphs[0].add_run(str(ht)); rn.bold = True; rn.font.size = Pt(fs)
        rn.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); shade(c, "0B3D6B")
    for _, row in df.iterrows():
        cs = t.add_row().cells
        for j, col in enumerate(df.columns):
            v = row[col]
            txt = f"{v*100:.1f} %" if (col in pct and isinstance(v, (int, float))) else \
                  (f"{v:.3f}" if isinstance(v, float) else str(v))
            cs[j].text = ""; rn = cs[j].paragraphs[0].add_run(txt); rn.font.size = Pt(fs)


def note(doc, t):
    p = doc.add_paragraph(); r = p.add_run(t); r.italic = True; r.font.size = Pt(9); r.font.color.rgb = RED
    p.paragraph_format.space_after = Pt(3)


def build():
    doc = Document(); setfont(doc)
    for s in doc.sections:
        s.top_margin = Cm(1.6); s.bottom_margin = Cm(1.6); s.left_margin = Cm(1.9); s.right_margin = Cm(1.9)

    head = csv("honest_headline_bssid.csv").iloc[0]

    # En-tête compact (pas de page de garde séparée pour tenir en 10 pages)
    para(doc, "UTT — IF23 P · Géolocalisation indoor", size=10, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "Localisation indoor par fingerprinting WiFi", size=17, bold=True, color=NAVY,
         align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "Binôme 6 : Molka GHADDAB · Maxime LE GAL · Virgile TUPET — enseignante : F. Chehade — juin 2026",
         size=9.5, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)

    # 1. Introduction
    h(doc, "1. Introduction et objectifs")
    para(doc,
        "Le but du projet est de localiser en temps réel un ordinateur à l'intérieur du bâtiment de "
        "l'UTT à partir des signaux WiFi qu'il capte. Deux tâches sont visées : (1) reconnaître la "
        "zone (salle) parmi plusieurs, par classification ; (2) estimer la position exacte (x, y) à "
        "l'intérieur d'une zone, par régression. Le principe est le fingerprinting : on enregistre "
        "à l'avance la signature radio (les puissances RSSI des points d'accès) de chaque lieu, "
        "puis on compare une nouvelle mesure à cette base.")
    para(doc,
        "Ce rapport couvre l'étude des données, la classification et son évaluation, l'évolution "
        "des performances dans le temps (jeux ALT1, ALT2, ALT3), la régression intra-zone, et "
        "plusieurs évolutions avancées (sélection de réseaux, modèle bayésien, optimisation, "
        "approche hiérarchique).")

    # 2. Données
    h(doc, "2. Données et collecte")
    para(doc,
        "Chaque mesure associe un réseau (SSID, BSSID), sa puissance (RSSI en dBm) et l'instant de "
        "mesure. Pour chaque point, on prend plusieurs scans que l'on moyenne par réseau, ce qui "
        "donne un vecteur de RSSI ; les réseaux non vus reçoivent -100 dBm. La base est organisée "
        "par bâtiment, étage, zone, coordonnées (x, y), réseau et instant.")
    para(doc,
        "La consigne prévoit trois campagnes successives, qui servent à mesurer la tenue du modèle "
        "dans le temps : ALT1 (30 points par zone), ALT2 (15 points) et ALT3 (5 points). Le jeu "
        "ALT1 utilisé ici compte 1 498 snapshots sur 24 zones (bâtiments P et S, plusieurs "
        "étages). Les zones n'ont pas toutes le même nombre de mesures, ce qui nous conduit à "
        "regarder le macro-F1 et la balanced accuracy en plus de l'accuracy.")
    diff = pd.DataFrame({
        "Jeu": ["ALT1", "ALT2", "ALT3"],
        "Période": ["février 2026", "28 avril 2026", "juin 2026 (à venir)"],
        "Zones": ["24 (sous-zones A/B)", "10 (salles entières)", "—"],
        "Snapshots": ["1 498", "452", "—"],
        "Rôle": ["apprentissage", "test temporel + ré-app.", "test final (stratégie 3)"],
    })
    table(doc, diff, list(diff.columns), fs=8.5)
    para(doc, "Tableau 1 — Les trois campagnes successives. ALT1 et ALT2 sont intégrées ; ALT3 "
              "reste à collecter (placeholder prêt).", italic=True, size=8.5, color=GREY)
    note(doc,
        "Les zones d'ALT2 sont des salles entières (P102, S103…) ; les sous-zones A/B d'ALT1 sont "
        "fusionnées pour la comparaison. 7 zones sont communes aux deux campagnes, 3 sont nouvelles "
        "en ALT2.")

    # 3. Étude des données ALT1
    h(doc, "3. Étude des données et sélection des réseaux")
    para(doc,
        "Le jeu ALT1 contient 104 points d'accès retenus (présents dans au moins 5 snapshots). "
        "Beaucoup sont redondants : on relève 285 paires de réseaux très corrélées (|r| > 0,9), "
        "principalement les réseaux institutionnels (UTT, eduroam, UTTetudiants) portés par les "
        "mêmes bornes physiques. Cette redondance justifie une sélection de réseaux plutôt que "
        "d'utiliser les 104 sans tri. La couverture est très inégale : quelques réseaux stables, "
        "et une longue traîne de réseaux éphémères (figure 1).")
    fig(doc, "fig17_couverture.png",
        "Figure 1 — Couverture des réseaux : quelques réseaux stables, beaucoup d'éphémères.", w=4.8)

    # 4. Classification
    h(doc, "4. Classification en zones")
    para(doc,
        "Nous avons comparé les modèles vus en cours (k plus proches voisins, bayésien, SVM à "
        "noyau RBF, forêts aléatoires, ExtraTrees, réseau de neurones), d'abord avec tous les "
        "réseaux puis avec une sélection. Les méthodes à base d'arbres arrivent en tête.")
    para(doc,
        "Point central : la façon d'évaluer change tout. Comme les snapshots d'une même zone "
        "viennent d'une seule session et se ressemblent beaucoup, un découpage aléatoire "
        "train/test donne un score trop optimiste (le modèle retrouve en test des mesures presque "
        "identiques à l'entraînement). En séparant les données dans le temps, on obtient une mesure "
        "réaliste. La figure 2 illustre cette descente : partant de 98 % au test aléatoire, on "
        "rejoint la performance réellement observée en utilisation directe (environ 70 à 80 %) en "
        "ajoutant les facteurs réalistes.")
    fig(doc, "fig13_escalier.png",
        "Figure 2 — Du test aléatoire (98 %) à la performance réelle (~75 %) en rendant "
        "l'évaluation réaliste.", w=5.4)
    sb = csv("eval_summary_bssid.csv")
    show = sb[["model", "A_aleatoire", "C2_groupkfold"]].copy()
    table(doc, show, ["Modèle", "Test aléatoire", "Évaluation honnête"],
          pct=["A_aleatoire", "C2_groupkfold"])
    para(doc, "Tableau 2 — Accuracy par modèle : test aléatoire vs évaluation honnête (BSSID).",
         italic=True, size=8.5, color=GREY)
    para(doc,
        f"Le meilleur modèle honnête est ExtraTrees, autour de 82 %. Même en évaluation honnête, la "
        f"bonne zone est dans les trois premières propositions dans {head['top3']*100:.1f} % des "
        "cas, ce qui est précieux pour proposer une courte liste de zones. L'analyse de "
        "l'importance des réseaux montre toutefois que le modèle s'appuie en partie sur des "
        "téléphones personnels présents lors de la collecte (figure 3) : utiles sur le moment, mais "
        "instables d'une campagne à l'autre. Un modèle limité aux réseaux institutionnels est moins "
        "performant mais plus fiable dans la durée.")
    fig(doc, "fig4_importance_reseaux.png",
        "Figure 3 — Réseaux les plus discriminants. Les deux premiers sont des téléphones personnels.", w=5.0)

    # 5. Évolution temporelle
    h(doc, "5. Évolution temporelle (ALT1 → ALT2 → ALT3)")
    ev = pd.read_csv(FIG / "evolution_real.csv")
    s1 = float(ev.loc[ev["strategie"].str.startswith("S1"), "accuracy"].iloc[0])
    s2 = float(ev.loc[ev["strategie"].str.startswith("S2"), "accuracy"].iloc[0])
    para(doc,
        "C'est le résultat le plus important du projet. On vérifie ici si un modèle entraîné en "
        "février tient toujours sur des mesures prises en avril (ALT2), trois mois plus tard. Trois "
        "stratégies sont prévues : S1, apprentissage sur ALT1 et test sur ALT2 ; S2, apprentissage "
        "sur ALT1 plus la moitié d'ALT2 et test sur le reste ; S3, apprentissage sur ALT1 et ALT2 "
        "et test sur ALT3 (à venir).")
    para(doc,
        f"En stratégie 1, l'accuracy chute à {s1*100:.0f} % : le modèle de février ne reconnaît "
        "plus la moitié des mesures d'avril. Certaines zones tombent même à zéro (figure 3), car "
        "leur environnement radio a changé (bornes déplacées ou remplacées). En réintégrant "
        f"seulement la moitié d'ALT2 à l'entraînement (stratégie 2), on remonte à {s2*100:.0f} %. "
        "Ce contraste mesure précisément le « drift » du WiFi : un modèle de fingerprinting doit "
        "être ré-entraîné régulièrement.")
    fig(doc, "fig18_evolution.png",
        "Figure 4 — Évolution ALT1 → ALT2. À gauche : accuracy par zone (stratégie 1), certaines "
        "zones s'effondrent. À droite : drift global puis récupération après ré-apprentissage.", w=6.0)
    note(doc, "Stratégie 3 (test sur ALT3) : en attente de la collecte ALT3 (juin). Le code est "
              "prêt (scripts/16) et le tableau se complétera automatiquement.")

    # 6. Régression
    h(doc, "6. Régression intra-zone (position exacte)")
    rr = pd.read_csv(FIG / "regression_results.csv").sort_values("MAE_m")
    best = rr.iloc[0]
    para(doc,
        "Pour estimer la position (x, y) à l'intérieur d'une zone, nous avons relevé 46 points de "
        "coordonnées connues dans une salle (repère local, zone d'environ 5 × 7 m) avec le RSSI de "
        "52 points d'accès. Plusieurs régresseurs multi-sorties prédisent (x, y) à partir du "
        "vecteur de RSSI ; la qualité se mesure par l'erreur euclidienne (moyenne, médiane) et la "
        "courbe cumulative des erreurs (CDF).")
    table(doc, rr.rename(columns={"modele": "Modèle", "MAE_m": "MAE (m)", "mediane_m": "Médiane (m)",
          "p90_m": "P90 (m)", "max_m": "Max (m)"}),
          ["Modèle", "MAE (m)", "Médiane (m)", "P90 (m)", "Max (m)"], fs=8.5)
    para(doc, "Tableau 3 — Erreur de position par modèle (validation croisée).",
         italic=True, size=8.5, color=GREY)
    para(doc,
        f"Le meilleur modèle ({best['modele']}) atteint une erreur médiane de {best['mediane_m']:.1f} m "
        f"et une moyenne de {best['MAE_m']:.1f} m, soit une position estimée à environ un mètre près "
        "dans une salle — un résultat correct vu la taille de la zone et le petit nombre de points. "
        "La figure 5 montre la CDF des erreurs et la figure 6 compare positions vraies et prédites.")
    fig(doc, "fig19_regression_cdf.png",
        "Figure 5 — CDF de l'erreur de position : médiane ≈ 1,2 m.", w=4.6)
    fig(doc, "fig20_regression_scatter.png",
        "Figure 6 — Positions vraies (vert) vs prédites (rouge) ; les traits relient chaque paire.", w=3.9)
    note(doc, "Limite : un seul local cartographié et 46 points. À étendre à d'autres zones.")

    # 7. Évolutions avancées
    h(doc, "7. Évolutions avancées")
    by = csv("bayesian_bssid.csv"); bc = by[by["modele"].str.contains("cours")].iloc[0]
    tu = csv("hyperparam_tuning_bssid.csv")
    et_gain = tu[tu["modele"] == "ExtraTrees"].iloc[0]
    hh = csv("hierarchical_bssid.csv")
    gfloor = float(hh.loc[hh["approche"].str.contains("étage"), "accuracy"].iloc[0])
    bullet(doc, f"nous l'avons programmé (RSSI gaussien par réseau, présence modélisée, décision au "
                f"maximum a posteriori). Il atteint {bc['groupkfold_acc']*100:.0f} % en évaluation "
                "honnête, comparable aux forêts.", "Modèle bayésien (cours) — ")
    bullet(doc, f"recherche sur grille avec validation honnête ; gains réels mais modestes "
                f"(ExtraTrees +{et_gain['gain']*100:.0f} pt). Le protocole et les réseaux comptent "
                "plus que le réglage.", "Optimisation des hyperparamètres — ")
    bullet(doc, f"prédire d'abord l'étage (fiable à {gfloor*100:.0f} %) puis la salle. N'améliore "
                "pas le score global ici, mais la fiabilité de l'étage est une base solide pour la "
                "cartographie.", "Approche hiérarchique (étage → salle) — ")
    bullet(doc, "un estimateur log-distance (exposant de propagation n ≈ 2,8, calibré sur un scan "
                "de référence) relie puissance et distance à une borne. Bonus hors consigne, utile "
                "au diagnostic.", "Estimation de distance — ")
    bullet(doc, "limiter les zones possibles d'un instant à l'autre à partir de la zone estimée et "
                "d'une vitesse maximale (graphe de transition). Piste prévue.", "Mobilité — ")
    fig(doc, "fig12_hierarchique.png",
        "Figure 7 — Approche hiérarchique : la détection d'étage (94 %) servira à la cartographie.", w=4.8)

    # 8. Conclusion
    h(doc, "8. Conclusion")
    para(doc,
        "Le projet traite les deux tâches demandées : la classification en zones et la régression "
        "de position (x, y), avec une erreur médiane d'environ 1,2 m. Sur le hold-out, la "
        "classification atteint des scores très élevés, mais le vrai enseignement vient de "
        "l'évolution temporelle : sur une nouvelle campagne (ALT2, trois mois plus tard), "
        "l'accuracy chute à environ 47 %, et ne se rétablit qu'en ré-intégrant des mesures "
        "récentes. Le fingerprinting WiFi fonctionne, mais il faut le ré-entraîner régulièrement "
        "car l'environnement radio change.")
    para(doc,
        "Il reste à collecter ALT3 (juin) pour compléter la stratégie 3 (le code est prêt), à "
        "étendre la régression à d'autres zones, et à exploiter les pistes avancées (hiérarchique, "
        "mobilité, cartographie). L'apport méthodologique du projet est d'avoir distingué un score "
        "de hold-out flatteur d'une performance réelle mesurée sur de nouvelles données.")

    OUT.parent.mkdir(parents=True, exist_ok=True); doc.save(OUT)
    print("Rapport condensé généré :", OUT)


if __name__ == "__main__":
    build()
