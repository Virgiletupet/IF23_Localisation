"""
06_generate_report.py
=====================
Génère le rapport Word (.docx) ALT3 du projet de localisation indoor WiFi.
Lit les CSV/figures de reports/figures/ et produit reports/Rapport_IF23_ALT3.docx.
Rédaction volontairement sobre (ton élève-ingénieur, première personne).
"""
from __future__ import annotations

import sys
from pathlib import Path

try:
    sys.stdout.reconfigure(encoding="utf-8")
except Exception:
    pass

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm, Inches

ROOT = Path(__file__).resolve().parents[1]
FIG = ROOT / "reports" / "figures"
OUT = ROOT / "reports" / "Rapport_IF23_ALT3.docx"

NAVY = RGBColor(0x0B, 0x3D, 0x6B)
RED = RGBColor(0xC0, 0x39, 0x2B)
GREY = RGBColor(0x55, 0x55, 0x55)


def csv(name):
    return pd.read_csv(FIG / name)


def setfont(doc):
    st = doc.styles["Normal"]
    st.font.name = "Calibri"; st.font.size = Pt(11)
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")


def h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = NAVY if level <= 2 else GREY
    return p


def para(doc, text, italic=False, size=11, align=None, color=None, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic; r.bold = bold; r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    if align: p.alignment = align
    return p


def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix); r.bold = True
    p.add_run(text)
    return p


def figure(doc, filename, caption, width_in=6.1):
    if not (FIG / filename).exists():
        para(doc, f"[figure manquante: {filename}]", italic=True, color=RED); return
    doc.add_picture(str(FIG / filename), width=Inches(width_in))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption); r.italic = True; r.font.size = Pt(9); r.font.color.rgb = GREY


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), hexcolor)
    tcPr.append(sh)


def table_from_df(doc, df, headers=None, pct_cols=None, highlight_row=None, fontsize=9):
    pct_cols = pct_cols or []
    headers = headers or list(df.columns)
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, htext in enumerate(headers):
        c = t.rows[0].cells[j]; c.text = ""
        run = c.paragraphs[0].add_run(str(htext)); run.bold = True; run.font.size = Pt(fontsize)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); shade(c, "0B3D6B")
    for _, row in df.iterrows():
        cells = t.add_row().cells
        for j, col in enumerate(df.columns):
            v = row[col]
            if col in pct_cols and isinstance(v, (int, float)):
                txt = f"{v*100:.1f} %"
            elif isinstance(v, float):
                txt = f"{v:.3f}"
            else:
                txt = str(v)
            cells[j].text = ""
            run = cells[j].paragraphs[0].add_run(txt); run.font.size = Pt(fontsize)
            if highlight_row is not None and row.get("model") == highlight_row:
                run.bold = True; shade(cells[j], "E8F0FA")
    return t


def note(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.italic = True; r.font.size = Pt(9.5); r.font.color.rgb = RED


# ====================================================================== build
def build():
    doc = Document(); setfont(doc)
    for s in doc.sections:
        s.top_margin = Cm(2.2); s.bottom_margin = Cm(2.2)
        s.left_margin = Cm(2.4); s.right_margin = Cm(2.4)

    # ---------------------------------------------------------- page de garde
    para(doc, "UNIVERSITÉ DE TECHNOLOGIE DE TROYES",
         align=WD_ALIGN_PARAGRAPH.CENTER, size=15, bold=True, color=NAVY)
    para(doc, "UV IF23 — Localisation et navigation", align=WD_ALIGN_PARAGRAPH.CENTER,
         size=12, color=GREY)
    for _ in range(3): doc.add_paragraph()
    para(doc, "Localisation indoor par fingerprinting WiFi",
         align=WD_ALIGN_PARAGRAPH.CENTER, size=22, bold=True, color=NAVY)
    para(doc, "Détection automatique de salle dans le bâtiment de l'UTT",
         align=WD_ALIGN_PARAGRAPH.CENTER, size=13, color=GREY)
    doc.add_paragraph()
    para(doc, "Rapport de projet — jalons ALT1, ALT2 et ALT3",
         align=WD_ALIGN_PARAGRAPH.CENTER, size=12, italic=True)
    for _ in range(5): doc.add_paragraph()
    para(doc, "Binôme 6 : Molka GHADDAB — Maxime LE GAL — Virgile TUPET",
         align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
    para(doc, "Enseignante : Farah Chehade — Semestre P26 (juin 2026)",
         align=WD_ALIGN_PARAGRAPH.CENTER, size=11, color=GREY)
    doc.add_page_break()

    head = csv("honest_headline_bssid.csv").iloc[0]

    # ---------------------------------------------------------- résumé
    h(doc, "Résumé", 1)
    para(doc,
        "Ce projet cherche à reconnaître automatiquement la salle où se trouve un utilisateur "
        "dans le bâtiment de l'UTT, en n'utilisant que les signaux WiFi captés par un téléphone. "
        "Nous avons travaillé sur trois jalons : un premier essai en février (ALT1), un pipeline "
        "plus complet en avril (ALT2), et le travail présenté ici en juin (ALT3).")
    para(doc,
        "Le rendu d'avril annonçait 99 % de bonne classification, et même 100 % sur un test "
        "présenté comme « live ». Ces chiffres ne collaient pas avec ce que nous observions en "
        "vrai, où le système se trompait régulièrement. En reprenant le code, nous avons compris "
        "pourquoi : la façon d'évaluer le modèle était trop optimiste. D'une part le découpage "
        "entre données d'entraînement et de test était aléatoire, alors que nos mesures se "
        "ressemblent énormément d'un instant au suivant ; d'autre part le test « live » se faisait "
        "en réalité sur les données d'entraînement.")
    para(doc,
        "Quand on évalue le modèle correctement, en séparant les données dans le temps, la "
        f"précision tombe autour de {head['accuracy']*100:.0f} à 82 %. Ce niveau correspond à ce "
        "que nous mesurions réellement en utilisation directe, et il est cohérent avec ce qu'on "
        "trouve dans la littérature sur le sujet (80 à 95 %). Autrement dit, le système marche "
        "à peu près, mais pas à 99 %.")
    para(doc,
        "En plus de ce diagnostic, ce jalon ajoute plusieurs choses : une comparaison honnête des "
        "modèles vus en cours (k plus proches voisins, bayésien, SVM, forêts aléatoires, réseau de "
        "neurones), une analyse qui montre que le modèle s'appuyait en partie sur les téléphones "
        "personnels des personnes ayant fait les mesures, un modèle plus propre qui s'en passe, et "
        "une première brique hiérarchique (étage puis salle) qui prépare la cartographie.")

    # ---------------------------------------------------------- 1. intro
    h(doc, "1. Introduction", 1)
    h(doc, "1.1 Contexte", 2)
    para(doc,
        "À l'intérieur d'un bâtiment, le GPS ne fonctionne pas : les murs atténuent le signal et "
        "il n'y a pas de notion d'étage. Le WiFi, lui, est présent partout, ce qui en fait une "
        "base intéressante pour se localiser sans installer de matériel supplémentaire. Les usages "
        "possibles vont du guidage de visiteurs à la gestion des salles ou à des services de "
        "campus.")
    h(doc, "1.2 Objectifs", 2)
    bullet(doc, "Mettre en place une collecte de mesures WiFi et un prétraitement adapté.")
    bullet(doc, "Comparer plusieurs algorithmes de classification vus en cours.")
    bullet(doc, "Évaluer honnêtement ce que le système sait réellement faire.")
    bullet(doc, "Comprendre l'écart entre les chiffres annoncés et le comportement en direct.")
    bullet(doc, "Proposer des pistes d'amélioration et préparer la cartographie.")
    h(doc, "1.3 Périmètre", 2)
    para(doc,
        "Le système couvre 24 zones du bâtiment : des salles de cours (S101A à S204B), des locaux "
        "(P101A à P103A), des paliers et des couloirs, répartis sur trois niveaux. Cette variété "
        "de lieux rend la tâche plus difficile qu'il n'y paraît, car certaines zones voisines se "
        "ressemblent beaucoup du point de vue radio.")

    # ---------------------------------------------------------- 2. état de l'art
    h(doc, "2. État de l'art", 1)
    h(doc, "2.1 Pourquoi le fingerprinting", 2)
    para(doc,
        "La puissance d'un signal WiFi diminue avec la distance, mais en intérieur les murs, les "
        "réflexions et les obstacles rendent cette relation trop imprévisible pour calculer une "
        "position par simple triangulation. Le fingerprinting contourne le problème : au lieu de "
        "modéliser la propagation, on enregistre à l'avance la « signature » radio de chaque zone, "
        "puis on compare une nouvelle mesure à cette base.")
    para(doc,
        "La méthode comporte deux phases. Pendant la phase de calibration (hors ligne), on "
        "parcourt chaque zone en relevant les puissances reçues (RSSI) de chaque point d'accès. "
        "Pendant la phase d'estimation (en ligne), on prend une nouvelle mesure et on cherche à "
        "quelle zone elle ressemble le plus. Le problème revient donc à une classification.")
    h(doc, "2.2 Les modèles vus en cours", 2)
    bullet(doc, "on compare la mesure aux empreintes connues (distance euclidienne) et on prend "
                "la zone majoritaire parmi les voisins. C'est la référence du fingerprinting.",
                "k plus proches voisins — ")
    bullet(doc, "on applique la règle de Bayes, P(zone|mesure) proportionnel à "
                "P(mesure|zone)·P(zone), en supposant les points d'accès indépendants et un RSSI "
                "gaussien par point d'accès.", "Bayésien — ")
    bullet(doc, "on cherche la frontière qui sépare le mieux les zones, avec un noyau RBF pour "
                "gérer les cas non linéaires.", "SVM — ")
    bullet(doc, "un ensemble d'arbres de décision ; robuste au bruit et capable d'indiquer quels "
                "points d'accès comptent le plus.", "Forêts aléatoires / ExtraTrees — ")
    bullet(doc, "un perceptron multicouche, capable d'apprendre des relations complexes mais plus "
                "gourmand en données.", "Réseau de neurones — ")
    h(doc, "2.3 Comment on mesure la qualité", 2)
    para(doc,
        "Pour un classifieur de zones, on regarde la matrice de confusion, l'accuracy (part de "
        "bonnes réponses), et la précision et le rappel par zone. Comme nos zones n'ont pas toutes "
        "le même nombre de mesures, on ajoute le macro-F1 et la balanced accuracy, qui donnent le "
        "même poids à chaque zone. La Top-k accuracy indique si la bonne zone fait partie des k "
        "plus probables. Un point souligné en cours nous a servi de fil rouge : tester un modèle "
        "sur ses propres données d'entraînement mesure sa mémoire, pas sa capacité à généraliser.")

    # ---------------------------------------------------------- 3. données
    h(doc, "3. Données et collecte", 1)
    para(doc,
        "Les mesures (BSSID, SSID, RSSI en dBm, horodatage) ont été relevées avec une application "
        "mobile, lors de deux sorties : le 10 février 2026 pour 16 zones, puis le 13 février pour "
        "8 zones supplémentaires. Dans chaque zone, on se déplaçait sur quelques points et on "
        "laissait tourner les scans.")
    para(doc,
        "Nous regroupons les mesures par instant : un « snapshot » rassemble tous les points "
        "d'accès visibles à un moment donné, transformés en un vecteur de caractéristiques. Au "
        "total nous avons 1 498 snapshots sur 24 zones. Le nombre de snapshots varie pas mal d'une "
        "zone à l'autre (figure 1), ce qui justifie de regarder le macro-F1 en plus de l'accuracy.")
    figure(doc, "fig6_snapshots_par_zone.png",
           "Figure 1 — Nombre de snapshots par zone. Les classes sont déséquilibrées.")
    note(doc,
        "Une limite à garder en tête : toutes les mesures datent de février. Nous n'avons donc "
        "pas pu vérifier si le modèle reste bon un autre jour ou à un autre moment. C'est le sujet "
        "des perspectives (campagnes d'avril et de juin).")

    # ---------------------------------------------------------- 4. pipeline
    h(doc, "4. Traitement des données et caractéristiques", 1)
    para(doc,
        "Le module robust_localization.py regroupe le traitement : nettoyage (mise en forme des "
        "identifiants, RSSI dans des bornes valides), construction des snapshots, puis "
        "vectorisation. Pour chaque réseau présent dans au moins 5 snapshots, on calcule quelques "
        "statistiques de RSSI (moyenne, maximum, écart-type, présence), auxquelles s'ajoutent des "
        "indicateurs globaux du scan (nombre de réseaux visibles, RSSI moyen, etc.). Quand un "
        "réseau n'est pas détecté, on lui attribue -100 dBm.")
    para(doc,
        "Nous avons testé deux niveaux de granularité. Le premier identifie chaque point d'accès "
        "physique par son BSSID (104 retenus, 529 caractéristiques) : c'est le plus riche, et "
        "celui du modèle principal. Le second regroupe les points d'accès par nom de réseau, le "
        "SSID (74 caractéristiques) : plus compact, et c'est celui qu'utilise l'interface live.")
    note(doc,
        "Au passage, une imprécision du rendu d'avril : il parlait de « 120 SSID », alors que le "
        "modèle à 99 % repose en fait sur les BSSID (les points d'accès physiques). Nous "
        "rétablissons la distinction ici.")

    # ---------------------------------------------------------- 5. méthodo
    h(doc, "5. La question de l'évaluation", 1)
    para(doc,
        "Ce qui nous intéresse, ce n'est pas que le modèle reconnaisse des mesures déjà vues, "
        "mais qu'il sache classer de nouvelles mesures. Or les snapshots d'une même zone viennent "
        "tous d'une seule session de collecte : deux scans pris à quelques secondes d'intervalle, "
        "au même endroit, sont presque identiques.")
    h(doc, "5.1 Le piège du découpage aléatoire", 2)
    para(doc,
        "Si on partage les données au hasard entre entraînement et test, ces scans quasi "
        "identiques se retrouvent des deux côtés. Le modèle « reconnaît » alors en test des points "
        "qu'il a déjà presque vus en entraînement, et son score grimpe artificiellement. C'est une "
        "fuite de données. Le test « live » du rendu d'avril allait plus loin encore, puisqu'il "
        "prédisait directement sur les fichiers d'entraînement.")
    h(doc, "5.2 Comment nous évaluons honnêtement", 2)
    bullet(doc, "pour chaque zone, on trie les snapshots par ordre chronologique et on garde les "
                "70 % du début pour l'entraînement, les 30 % de la fin pour le test. Les deux "
                "ensembles sont ainsi séparés dans le temps.", "Découpage temporel (70/30) — ")
    bullet(doc, "on coupe chaque zone en 5 blocs successifs dans le temps, et la validation "
                "croisée garde toujours un bloc entier d'un seul côté. Cela donne une moyenne plus "
                "stable.", "Validation croisée par blocs (GroupKFold) — ")
    para(doc,
        "Nous comparons ces deux protocoles au découpage aléatoire et au test sur les données "
        "d'entraînement, en gardant exactement les mêmes caractéristiques, pour voir ce que change "
        "la seule façon d'évaluer.")
    h(doc, "5.3 Du test de base à la réalité du terrain", 2)
    sta = csv("staircase.csv")
    para(doc,
        "Pour relier les chiffres à notre expérience réelle, nous sommes partis du test de base "
        "(découpage aléatoire) et nous avons ajouté un à un les éléments qui rapprochent "
        "l'évaluation des conditions réelles : séparation dans le temps, puis découpage temporel "
        "strict, puis retrait des réseaux personnels, puis bruit de mesure. La figure 2 montre le "
        "résultat. On part de 98 % et, marche après marche, on retombe autour de 65 à 80 %, "
        "c'est-à-dire dans la zone de performance que nous observions réellement en utilisation "
        "directe (environ 70 à 80 %).")
    figure(doc, "fig13_escalier.png",
           "Figure 2 — Du test de base à la réalité. Chaque facteur réaliste ajouté fait baisser "
           "le score, qui rejoint la performance observée en utilisation directe.")
    para(doc,
        "C'est le point important du jalon : une fois l'évaluation rendue réaliste, le chiffre "
        "obtenu (autour de 75 %) correspond à ce que donne le système en vrai. Le 99 % n'était "
        "donc pas le vrai niveau du système, mais le résultat d'un test trop facile. La figure 3 "
        "complète ce constat en montrant que la précision se dégrade progressivement quand on "
        "ajoute du bruit sur le RSSI, sans chute brutale.")
    figure(doc, "fig14_bruit.png",
           "Figure 3 — Effet d'un bruit ajouté sur le RSSI au moment du test. La dégradation est "
           "progressive.")

    # ---------------------------------------------------------- 6. résultats
    h(doc, "6. Résultats", 1)
    h(doc, "6.1 L'écart entre test biaisé et évaluation honnête", 2)
    para(doc,
        "La figure 4 reprend, pour le meilleur modèle (ExtraTrees, BSSID), le score selon le "
        "protocole. On retrouve bien le 98,7 % du découpage aléatoire et le 100 % du test sur les "
        "données d'entraînement, puis la baisse une fois l'évaluation honnête. L'écart, de l'ordre "
        "de 20 points, est ce qui explique la différence avec le comportement en direct.")
    figure(doc, "fig1_protocoles_extratrees.png",
           "Figure 4 — Score selon le protocole d'évaluation (ExtraTrees, BSSID).")
    h(doc, "6.2 Comparaison des modèles", 2)
    sb = csv("eval_summary_bssid.csv")
    show = sb[["model", "A_aleatoire", "B_insample", "C1_holdout_temporel", "C2_groupkfold"]].copy()
    table_from_df(doc, show,
                  headers=["Modèle", "Aléatoire", "Sur l'entraînement", "Holdout temporel", "GroupKFold"],
                  pct_cols=["A_aleatoire", "B_insample", "C1_holdout_temporel", "C2_groupkfold"],
                  highlight_row="ExtraTrees")
    para(doc, "Tableau 1 — Accuracy par modèle et par protocole (BSSID). Les forêts restent "
              "devant, et le classement ne change pas en évaluation honnête.",
         italic=True, size=9, color=GREY)
    para(doc,
        "Les méthodes à base d'arbres (ExtraTrees, forêt aléatoire) arrivent en tête, suivies du "
        "SVM et du k-NN. Le bayésien gaussien simple est en retrait, mais nous verrons plus loin "
        "qu'une version mieux adaptée le rend compétitif. Surtout, tous les modèles perdent "
        "beaucoup en passant à l'évaluation honnête : le problème venait bien du protocole, pas "
        "d'un modèle en particulier.")
    figure(doc, "fig2_modeles_honnete.png",
           "Figure 5 — Comparaison des modèles en évaluation honnête (GroupKFold).")
    figure(doc, "fig3_biaise_vs_honnete.png",
           "Figure 6 — Écart entre test aléatoire et évaluation honnête, modèle par modèle.")
    h(doc, "6.3 Résultats par zone", 2)
    para(doc,
        "En évaluation honnête, le résultat est très inégal selon la zone (figure 7). Certaines "
        "zones bien isolées (les paliers, S104) restent parfaitement reconnues, tandis que "
        "d'autres s'effondrent (S101B, S203A) car elles se confondent avec une zone voisine. Ces "
        "confusions étaient totalement masquées par le test aléatoire.")
    figure(doc, "fig5_f1_par_zone.png",
           "Figure 7 — F1-score par zone en évaluation honnête.")
    figure(doc, "honest_confusion_matrix_bssid.png",
           "Figure 8 — Matrice de confusion honnête. Les confusions entre zones voisines "
           "réapparaissent.", width_in=5.8)
    h(doc, "6.4 BSSID ou SSID", 2)
    figure(doc, "fig8_bssid_vs_ssid.png",
           "Figure 9 — BSSID contre SSID en évaluation honnête. Regrouper par nom de réseau "
           "généralise un peu mieux.")
    h(doc, "6.5 Sur quoi le modèle s'appuie", 2)
    para(doc,
        "En regardant quels réseaux pèsent le plus dans la décision (figure 10), nous avons eu une "
        "surprise : les deux plus importants sont des téléphones personnels des personnes ayant "
        "fait la collecte (« s24 ultra de killian », « iphone de virgile »). Le modèle apprend "
        "donc en partie qui a mesuré où, et pas seulement les caractéristiques du lieu. Ces "
        "réseaux disparaissent ou changent d'une session à l'autre, ce qui contribue à la baisse "
        "de performance en utilisation réelle.")
    figure(doc, "fig4_importance_reseaux.png",
           "Figure 10 — Les 15 réseaux les plus discriminants. Les deux premiers sont des "
           "téléphones personnels.")
    h(doc, "6.6 Un modèle qui se passe des réseaux personnels", 2)
    sc = csv("sanitized_comparison.csv")
    nice = sc[["jeu_features", "n_features", "holdout_temp_acc", "groupkfold_acc"]].copy()
    nice["jeu_features"] = nice["jeu_features"].map(
        {"all": "Tous réseaux", "institution": "Réseaux institutionnels", "perso": "Réseaux personnels"})
    table_from_df(doc, nice,
                  headers=["Jeu de réseaux", "Nb caractéristiques", "Holdout temporel", "GroupKFold"],
                  pct_cols=["holdout_temp_acc", "groupkfold_acc"])
    para(doc, "Tableau 2 — Effet du retrait des réseaux personnels (ExtraTrees, BSSID, honnête).",
         italic=True, size=9, color=GREY)
    para(doc,
        "On observe deux choses. D'un côté, les réseaux personnels pris seuls font même un peu "
        "mieux que les réseaux institutionnels seuls : ils sont effectivement très discriminants, "
        "mais ils ne dureront pas. De l'autre, le modèle limité aux réseaux institutionnels reste "
        "correct et constitue une base plus fiable pour une vraie utilisation, car il ne dépend "
        "pas des téléphones présents le jour de la collecte.")
    figure(doc, "fig7_assaini.png",
           "Figure 11 — Comparaison : tous les réseaux, institutionnels seuls, personnels seuls.")
    h(doc, "6.7 Le modèle bayésien du cours", 2)
    by = csv("bayesian_bssid.csv")
    bc = by[by["modele"].str.contains("cours")].iloc[0]
    para(doc,
        "Nous avons programmé nous-mêmes le modèle bayésien vu en cours : une probabilité a priori "
        "par zone, un RSSI gaussien par point d'accès, et une décision au maximum a posteriori. "
        "Plutôt que de remplir les réseaux non détectés par -100 (ce que fait le bayésien standard "
        "et qui crée des distributions dégénérées), nous modélisons explicitement la présence ou "
        "l'absence de chaque point d'accès. Ce détail change beaucoup : la version soignée atteint "
        f"{bc['groupkfold_acc']*100:.0f} % en évaluation honnête, contre nettement moins pour le "
        "bayésien brut (figure 12). Elle devient ainsi comparable aux forêts.")
    figure(doc, "fig10_bayesien.png",
           "Figure 12 — Le bayésien soigné (présence modélisée) contre le bayésien brut, BSSID.")
    h(doc, "6.8 Réglage des hyperparamètres", 2)
    tu = csv("hyperparam_tuning_bssid.csv")
    para(doc,
        "Nous avons réglé les principaux hyperparamètres par recherche sur grille, avec la même "
        "validation croisée par blocs pour rester honnêtes. Les meilleurs réglages sont assez "
        "classiques (k=3 avec pondération par la distance pour le k-NN, un C élevé pour le SVM, "
        "des arbres un peu plus profonds pour ExtraTrees). Les gains restent faibles, de 1 à 3 "
        "points (figure 13 et tableau 3) : sur ce problème, le protocole et les caractéristiques "
        "comptent plus que le réglage fin.")
    table_from_df(doc, tu[["modele", "defaut", "optimise", "gain"]].copy(),
                  headers=["Modèle", "Par défaut", "Optimisé", "Gain"],
                  pct_cols=["defaut", "optimise", "gain"])
    para(doc, "Tableau 3 — Effet du réglage (accuracy GroupKFold, BSSID). ExtraTrees optimisé "
              "donne le meilleur score honnête du projet.", italic=True, size=9, color=GREY)
    figure(doc, "fig11_tuning.png",
           "Figure 13 — Gain du réglage des hyperparamètres, par modèle.")
    h(doc, "6.9 Une première approche hiérarchique (étage puis salle)", 2)
    hh = csv("hierarchical_bssid.csv")
    g = lambda lbl: float(hh.loc[hh["approche"].str.contains(lbl), "accuracy"].iloc[0])
    para(doc,
        "Le cours suggère de localiser d'abord grossièrement, puis d'affiner. Nous avons donc "
        "essayé de prédire d'abord l'étage, puis la salle au sein de cet étage. La prédiction de "
        f"l'étage est très fiable ({g('étage')*100:.0f} %). En revanche, enchaîner étage puis "
        f"salle ({g('étage prédit')*100:.0f} %) ne fait pas mieux que prédire directement la salle "
        f"({g('Plat')*100:.0f} %), car les erreurs d'étage se reportent ensuite. Si l'étage était "
        f"connu, on monterait à {g('oracle')*100:.0f} %. Cette fiabilité de l'étage reste utile : "
        "savoir avec certitude à quel niveau se trouve l'utilisateur est déjà une information "
        "concrète pour un plan interactif.")
    figure(doc, "fig12_hierarchique.png",
           "Figure 14 — Approche hiérarchique. La détection d'étage (94 %) servira de base à la "
           "cartographie.")

    # ---------------------------------------------------------- 7. jalons
    h(doc, "7. Les trois jalons", 1)
    a1 = csv("alt1_reproduction.csv").iloc[0]
    df3 = pd.DataFrame({
        "Aspect": ["Période", "Zones", "Identifiants", "Caractéristiques", "Meilleur modèle",
                   "Score annoncé", "Score réel (honnête)", "Bilan"],
        "ALT1": ["février", "16", "6 BSSID", "6", "Forêt aléatoire",
                 f"≈ {a1['biaise_aleatoire']*100:.0f} %",
                 f"≈ {a1['honnete_groupkfold']*100:.0f} %", "Premier essai"],
        "ALT2": ["avril", "24", "104 BSSID", "529", "ExtraTrees", "99 % / 100 % « live »",
                 "≈ 72–82 %", "Évaluation trop optimiste"],
        "ALT3": ["juin", "24", "104 BSSID (+ propre)", "529", "ExtraTrees / bayésien",
                 "—", "≈ 83 % au mieux", "Audit + nouveaux modèles"],
    })
    table_from_df(doc, df3, headers=list(df3.columns), fontsize=8.5)
    para(doc, "Tableau 4 — Les trois jalons. ALT3 ne contredit pas ALT2 : il en donne la vraie "
              "mesure et l'enrichit.", italic=True, size=9, color=GREY)
    para(doc,
        "Vu honnêtement, le projet a quand même bien progressé. En passant de 6 à 104 points "
        f"d'accès, la performance réelle est montée d'environ {a1['honnete_groupkfold']*100:.0f} % "
        "(ALT1) à environ 80 % (ALT2 et ALT3). Le travail d'avril était donc utile ; c'est "
        "seulement le chiffre annoncé (99 %) qui était trompeur. Ce jalon remet les chiffres au "
        "bon niveau et ajoute de nouveaux modèles ainsi qu'un réglage propre.")
    figure(doc, "fig9_progression_honnete.png",
           "Figure 15 — Progression réelle entre jalons, en évaluation honnête.")

    # ---------------------------------------------------------- 8. discussion
    h(doc, "8. Discussion", 1)
    h(doc, "8.1 Pourquoi 99 % et pas en vrai", 2)
    para(doc,
        "Trois raisons se cumulent. D'abord, les snapshots d'une même zone se ressemblent trop, et "
        "le découpage aléatoire en profite. Ensuite, le test « live » se faisait sur les données "
        "d'entraînement. Enfin, le modèle s'appuyait sur des réseaux personnels qui ne sont pas là "
        "le lendemain. En conditions réelles, ces trois appuis disparaissent, d'où la baisse.")
    h(doc, "8.2 Un résultat finalement crédible", 2)
    para(doc,
        "Autour de 80 % en validation honnête, le système se situe dans la fourchette de la "
        "littérature pour de la classification par pièce. C'est un résultat correct et défendable, "
        "bien plus crédible que le 99 % de départ, qui n'était pas réaliste pour un simple "
        "téléphone sans matériel dédié.")
    h(doc, "8.3 L'intérêt du Top-3", 2)
    para(doc,
        f"Même en évaluation honnête, la bonne zone est dans les trois premières propositions dans "
        f"{head['top3']*100:.1f} % des cas. Le système est donc fiable pour proposer une courte "
        "liste de zones probables. C'est exactement ce dont on a besoin pour la suite : proposer "
        "la zone, puis affiner la position sur un plan.")
    h(doc, "8.4 Limites", 2)
    bullet(doc, "Toutes les mesures datent de février : la tenue dans le temps n'est pas vérifiée.")
    bullet(doc, "Le modèle dépend encore un peu de réseaux non institutionnels.")
    bullet(doc, "Les zones voisines sont difficiles à départager, et les classes sont déséquilibrées.")
    bullet(doc, "Le modèle BSSID est assez lourd (environ 24 Mo) pour un usage embarqué.")

    # ---------------------------------------------------------- 9. perspectives
    h(doc, "9. Pistes d'amélioration", 1)
    bullet(doc, "refaire des campagnes à d'autres moments (avril, juin, différents horaires), puis "
                "entraîner sur février et tester sur avril. C'est le vrai test de robustesse.",
                "Tenir dans le temps — ")
    bullet(doc, "garder surtout les réseaux institutionnels, stables, et écarter les hotspots "
                "personnels.", "Choisir les bons réseaux — ")
    bullet(doc, "combiner le WiFi avec les capteurs du téléphone (accéléromètre, gyroscope) pour "
                "suivre le déplacement entre deux scans.", "Fusionner les capteurs — ")
    bullet(doc, "prédire la zone (ce travail), puis estimer la position à l'intérieur et la "
                "reporter sur un plan. C'est l'objet du jalon de cartographie.",
                "Vers la cartographie — ")

    # ---------------------------------------------------------- 10. conclusion
    h(doc, "10. Conclusion", 1)
    para(doc,
        "Au final, nous avons un système qui reconnaît la salle dans 24 zones du bâtiment, avec "
        "une performance honnête d'environ 80 %. Le principal apport de ce jalon n'est pas un "
        "score record, mais une mise au point : nous avons compris pourquoi le 99 % annoncé ne se "
        "retrouvait pas en vrai, nous l'avons démontré, et nous avons réévalué le système "
        "correctement. Le chiffre obtenu colle à ce que nous observions en utilisation directe.")
    para(doc,
        "Les pistes pour la suite sont claires : tester la tenue dans le temps, s'appuyer sur des "
        "réseaux plus stables, fusionner les capteurs, et passer à la cartographie en affinant la "
        "position à l'intérieur de la zone.")

    # ---------------------------------------------------------- 11. références
    h(doc, "11. Références", 1)
    refs = [
        "Bahl, P., & Padmanabhan, V. N. (2000). RADAR: An in-building RF-based user location and tracking system. IEEE INFOCOM.",
        "Youssef, M., & Agrawala, A. (2005). The Horus WLAN location determination system. ACM MobiSys.",
        "Geurts, P., Ernst, D., & Wehenkel, L. (2006). Extremely randomized trees. Machine Learning, 63(1).",
        "Breiman, L. (2001). Random forests. Machine Learning, 45(1).",
        "Torres-Sospedra, J., et al. (2014). UJIIndoorLoc: A multi-building/floor WLAN fingerprint database. IEEE IPIN.",
        "He, S., & Chan, S. H. G. (2016). Wi-Fi fingerprint-based indoor positioning: recent advances. IEEE Comm. Surveys & Tutorials, 18(1).",
        "Pedregosa, F., et al. (2011). Scikit-learn: Machine Learning in Python. JMLR, 12.",
        "Alshamaa, D. Supports de cours IF23 — Localisation indoor. UTT.",
    ]
    for i, r in enumerate(refs, 1):
        p = doc.add_paragraph(); p.add_run(f"[{i}] ").bold = True; p.add_run(r).font.size = Pt(10)

    # ---------------------------------------------------------- annexe
    doc.add_page_break()
    h(doc, "Annexe — Reproductibilité", 1)
    para(doc, "Le projet est organisé pour pouvoir être relancé facilement :")
    for line in [
        "src/robust_localization.py — caractéristiques et modèles.",
        "src/eval_utils.py — protocoles d'évaluation honnête.",
        "src/bayesian_localization.py — modèle bayésien du cours.",
        "scripts/02 à 12 — évaluation, modèle propre, ALT1, bayésien, réglage, hiérarchique, escalier.",
        "reports/figures/ — figures et résultats au format CSV.",
    ]:
        bullet(doc, line)
    para(doc, "Environnement : Python 3.10, scikit-learn 1.5.", size=10, color=GREY)
    h2 = csv("honest_headline_bssid.csv").iloc[0]
    para(doc,
        f"Repère chiffré (ExtraTrees, BSSID, holdout temporel) : accuracy {h2['accuracy']*100:.1f} %, "
        f"balanced accuracy {h2['balanced_accuracy']*100:.1f} %, macro-F1 {h2['macro_f1']*100:.1f} %, "
        f"Top-3 {h2['top3']*100:.1f} % ({int(h2['n_train'])} snapshots d'entraînement, "
        f"{int(h2['n_test'])} de test).", size=10)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print("Rapport généré :", OUT)


if __name__ == "__main__":
    build()
